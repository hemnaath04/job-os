"""Extract a JSON Resume from an uploaded PDF or DOCX using Claude.

For PDFs we use Claude's native document support — the file is sent inline
as a base64 `document` content block, so layout context is preserved and the
extractor sees citations, dates, and structure exactly as they appear.

For DOCX we extract text with python-docx first, then hand the text to
Claude. We never invent fields — empty inputs return empty arrays.
"""
from __future__ import annotations

import base64
import io
import json
from typing import Any

import anthropic
import structlog
from docx import Document
from pydantic import ValidationError

from job_os.settings import get_settings

log = structlog.get_logger(__name__)


SYSTEM_PROMPT = """You are a precise resume parser. Given a resume (PDF or text),
extract a JSON Resume document strictly matching the schema at
https://jsonresume.org. Return ONLY valid JSON, no commentary.

Rules:
- Use exact phrasing from the source — do not paraphrase bullets, summaries,
  or descriptions.
- If a field is not present in the source, leave it null or omit it. NEVER
  invent values, dates, locations, or metrics.
- Dates: prefer YYYY-MM format. Use YYYY if month is unknown. Omit endDate
  for ongoing entries (don't write "Present" or "Now").
- Keep highlights as separate array entries — one bullet per element.
- Group skills the way the source groups them; preserve category names verbatim.
"""

USER_PROMPT_PDF = """Extract this resume into JSON Resume format.

Respond with a single JSON object — no markdown fences, no prose.
"""

USER_PROMPT_TEXT = """Below is the text of a resume. Extract it into JSON Resume format.

<resume>
{text}
</resume>

Respond with a single JSON object — no markdown fences, no prose.
"""


def _strip_json_fence(text: str) -> str:
    t = text.strip()
    if t.startswith("```"):
        nl = t.find("\n")
        t = t[nl + 1 :] if nl != -1 else t
        if t.endswith("```"):
            t = t[:-3]
    return t.strip()


async def extract_json_resume_from_pdf(pdf_bytes: bytes) -> dict[str, Any]:
    s = get_settings()
    if not s.anthropic_api_key:
        raise RuntimeError("ANTHROPIC_API_KEY is not set; cannot extract from PDF.")

    client = anthropic.AsyncAnthropic(
        auth_token=s.anthropic_api_key,
        base_url=s.anthropic_base_url or None,
    )

    b64 = base64.standard_b64encode(pdf_bytes).decode("ascii")
    msg = await client.messages.create(
        model=s.anthropic_model_extract,
        max_tokens=8192,
        system=SYSTEM_PROMPT,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {"type": "base64", "media_type": "application/pdf", "data": b64},
                    },
                    {"type": "text", "text": USER_PROMPT_PDF},
                ],
            }
        ],
    )
    text = "".join(b.text for b in msg.content if b.type == "text")
    return json.loads(_strip_json_fence(text))


def _docx_to_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


async def extract_json_resume_from_docx(docx_bytes: bytes) -> dict[str, Any]:
    text = _docx_to_text(docx_bytes)
    return await extract_json_resume_from_text(text)


async def extract_json_resume_from_text(text: str) -> dict[str, Any]:
    s = get_settings()
    if not s.anthropic_api_key:
        raise RuntimeError("ANTHROPIC_API_KEY is not set; cannot extract from text.")

    client = anthropic.AsyncAnthropic(
        auth_token=s.anthropic_api_key,
        base_url=s.anthropic_base_url or None,
    )

    msg = await client.messages.create(
        model=s.anthropic_model_extract,
        max_tokens=8192,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": USER_PROMPT_TEXT.format(text=text[:50000])}],
    )
    body = "".join(b.text for b in msg.content if b.type == "text")
    return json.loads(_strip_json_fence(body))
